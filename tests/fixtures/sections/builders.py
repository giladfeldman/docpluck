"""Synthetic PDF/DOCX/HTML fixtures for section-identification tests.

Each builder emits a minimal document containing exactly the section
labels the test needs. Built on demand (not committed binaries) so tests
remain hermetic and the repo stays small.
"""

from __future__ import annotations

import io


def build_apa_single_study_pdf() -> bytes:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    y = 720
    def heading(t):
        nonlocal y
        c.setFont("Helvetica-Bold", 14); c.drawString(72, y, t); y -= 20
    def para(t):
        nonlocal y
        c.setFont("Helvetica", 11); c.drawString(72, y, t); y -= 20
    heading("Abstract"); para("This paper investigates X.")
    heading("Introduction"); para("Intro text.")
    heading("Methods"); para("We did things.")
    heading("Results"); para("We found stuff.")
    heading("Discussion"); para("It was great.")
    heading("References"); para("[1] Doe, J. (2020).")
    c.showPage(); c.save()
    return buf.getvalue()


def build_apa_multi_study_pdf() -> bytes:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    y = 720
    def heading(t):
        nonlocal y
        c.setFont("Helvetica-Bold", 14); c.drawString(72, y, t); y -= 20
    def para(t):
        nonlocal y
        c.setFont("Helvetica", 11); c.drawString(72, y, t); y -= 20
    heading("Abstract"); para("Multi-study paper.")
    heading("Introduction"); para("Intro text.")
    heading("Methods"); para("Study 1 methods.")
    heading("Results"); para("Study 1 results.")
    heading("Methods"); para("Study 2 methods.")
    heading("Results"); para("Study 2 results.")
    heading("General Discussion"); para("Synthesis.")
    heading("References"); para("[1] Doe, J. (2020).")
    c.showPage(); c.save()
    return buf.getvalue()


def build_html_with_real_headings() -> bytes:
    return (
        b"<html><body>"
        b"<h2>Abstract</h2><p>This paper investigates X.</p>"
        b"<h2>Methods</h2><p>We did things.</p>"
        b"<h2>References</h2><p>[1] Doe, J. (2020).</p>"
        b"</body></html>"
    )


def build_docx_with_real_headings() -> bytes:
    from docx import Document
    d = Document()
    d.add_heading("Abstract", level=2); d.add_paragraph("This paper investigates X.")
    d.add_heading("Methods", level=2); d.add_paragraph("We did things.")
    d.add_heading("References", level=2); d.add_paragraph("[1] Doe, J. (2020).")
    buf = io.BytesIO(); d.save(buf)
    return buf.getvalue()
