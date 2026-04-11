"""Tests for docpluck.extract_docx — DOCX text extraction via mammoth.

Tests use python-docx to programmatically create DOCX fixtures so they run
without external test files. The real end-to-end validation is in the
benchmark test suite (test_benchmark_docx_html.py).
"""
import io

import pytest

pytest.importorskip("mammoth", reason="mammoth not installed (pip install docpluck[docx])")
pytest.importorskip("bs4", reason="beautifulsoup4 not installed (pip install docpluck[html])")
pytest.importorskip("docx", reason="python-docx not installed (dev dependency)")

from docx import Document
from docpluck import extract_docx, normalize_text, NormalizationLevel


def _make_docx(paragraphs: list[str]) -> bytes:
    """Create an in-memory DOCX with the given paragraphs and return bytes."""
    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_heading(title: str, body: list[str]) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for para in body:
        doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Basic extraction
# ---------------------------------------------------------------------------

class TestBasicExtraction:
    def test_returns_tuple(self):
        docx = _make_docx(["Hello world"])
        result = extract_docx(docx)
        assert isinstance(result, tuple)
        assert len(result) == 2

    def test_method_is_mammoth(self):
        docx = _make_docx(["Hello world"])
        _, method = extract_docx(docx)
        assert method == "mammoth"

    def test_single_paragraph(self):
        docx = _make_docx(["Hello world"])
        text, _ = extract_docx(docx)
        assert "Hello world" in text

    def test_multiple_paragraphs(self):
        docx = _make_docx(["First paragraph", "Second paragraph", "Third paragraph"])
        text, _ = extract_docx(docx)
        assert "First paragraph" in text
        assert "Second paragraph" in text
        assert "Third paragraph" in text

    def test_paragraphs_separated_by_newline(self):
        docx = _make_docx(["Alpha", "Beta"])
        text, _ = extract_docx(docx)
        assert "AlphaBeta" not in text
        assert "\n" in text

    def test_empty_document(self):
        docx = _make_docx([])
        text, _ = extract_docx(docx)
        # Empty doc is still valid; should return a string (possibly empty)
        assert isinstance(text, str)


# ---------------------------------------------------------------------------
# Headings and structure
# ---------------------------------------------------------------------------

class TestStructure:
    def test_heading_extracted(self):
        docx = _make_docx_with_heading("Study Results", ["The effect was significant."])
        text, _ = extract_docx(docx)
        assert "Study Results" in text
        assert "The effect was significant." in text

    def test_heading_separated_from_body(self):
        docx = _make_docx_with_heading("Title", ["Body"])
        text, _ = extract_docx(docx)
        assert "TitleBody" not in text


# ---------------------------------------------------------------------------
# Unicode and academic content
# ---------------------------------------------------------------------------

class TestAcademicContent:
    def test_smart_quotes_preserved(self):
        """Smart quotes should come through; normalization handles conversion."""
        docx = _make_docx(["He said \u201chello\u201d to her."])
        text, _ = extract_docx(docx)
        assert "\u201c" in text or '"' in text
        assert "hello" in text

    def test_smart_quotes_normalize_to_ascii(self):
        """End-to-end: smart quotes in DOCX should normalize to ASCII quotes."""
        docx = _make_docx(["He said \u201chello\u201d to her."])
        text, _ = extract_docx(docx)
        normalized, _ = normalize_text(text, NormalizationLevel.standard)
        assert '"hello"' in normalized

    def test_unicode_characters(self):
        docx = _make_docx(["café naïve résumé"])
        text, _ = extract_docx(docx)
        assert "café" in text

    def test_statistical_values(self):
        docx = _make_docx([
            "Results",
            "We found a significant effect, F(1, 30) = 4.42, p = .043.",
            "The correlation was r(261) = -0.73, 95% CI [-0.78, -0.67], p < .001.",
        ])
        text, _ = extract_docx(docx)
        assert "F(1, 30) = 4.42" in text
        assert "r(261) = -0.73" in text
        assert "95% CI [-0.78, -0.67]" in text

    def test_ligatures_can_be_normalized(self):
        """DOCX with fi/fl ligatures (uncommon in raw DOCX but possible)."""
        docx = _make_docx(["The e\ufb03cient operation was signi\ufb01cant."])
        text, _ = extract_docx(docx)
        normalized, _ = normalize_text(text, NormalizationLevel.standard)
        assert "efficient" in normalized
        assert "significant" in normalized


# ---------------------------------------------------------------------------
# Error handling
# ---------------------------------------------------------------------------

class TestErrors:
    def test_invalid_bytes_raises(self):
        with pytest.raises(Exception):
            extract_docx(b"not a docx file")

    def test_empty_bytes_raises(self):
        with pytest.raises(Exception):
            extract_docx(b"")

    def test_pdf_bytes_raise(self):
        # Feeding a PDF header to DOCX extractor should fail cleanly
        with pytest.raises(Exception):
            extract_docx(b"%PDF-1.4\n...")


# ---------------------------------------------------------------------------
# Integration with normalization and quality scoring
# ---------------------------------------------------------------------------

class TestIntegration:
    def test_normalization_pipeline(self):
        docx = _make_docx([
            "The results showed a significant effect,",
            "F(1, 30) = 4.42, p = .043.",
            "Mean difference was M = 2.34, SD = 0.89.",
        ])
        text, _ = extract_docx(docx)
        normalized, report = normalize_text(text, NormalizationLevel.academic)
        assert isinstance(normalized, str)
        assert len(report.steps_applied) > 0
        assert "4.42" in normalized

    def test_quality_scoring(self):
        from docpluck import compute_quality_score
        docx = _make_docx([
            "The study examined the effect of treatment on cognitive performance.",
            "Results indicated a significant main effect of condition.",
            "The data were analyzed using a mixed-effects model with random intercepts.",
        ])
        text, _ = extract_docx(docx)
        quality = compute_quality_score(text)
        assert quality["score"] >= 80  # Clean DOCX text should score high
        assert quality["confidence"] == "high"
