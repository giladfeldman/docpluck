"""Filter sugar: extract_pdf(sections=...) returns concatenated section text."""

import io

import pytest

pytest.importorskip("pdfplumber")
pytest.importorskip("reportlab")


def _make_pdf() -> bytes:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.setFont("Helvetica-Bold", 14); c.drawString(72, 720, "Abstract")
    c.setFont("Helvetica", 11); c.drawString(72, 700, "Abstract body.")
    c.setFont("Helvetica-Bold", 14); c.drawString(72, 660, "Methods")
    c.setFont("Helvetica", 11); c.drawString(72, 640, "Methods body.")
    c.setFont("Helvetica-Bold", 14); c.drawString(72, 600, "References")
    c.setFont("Helvetica", 11); c.drawString(72, 580, "[1] Doe.")
    c.showPage(); c.save()
    return buf.getvalue()


def test_no_filter_preserves_byte_identical_output():
    """When sections= is None (default), output equals current extract_pdf()."""
    from docpluck import extract_pdf
    text_a, _ = extract_pdf(_make_pdf())
    text_b, _ = extract_pdf(_make_pdf(), sections=None)
    assert text_a == text_b


def test_filter_returns_only_requested_sections():
    from docpluck import extract_pdf
    text, _ = extract_pdf(_make_pdf(), sections=["abstract", "references"])
    assert "Abstract body" in text
    assert "[1] Doe" in text
    assert "Methods body" not in text


def test_filter_parity_with_sectioned_document():
    from docpluck import extract_pdf, extract_sections
    pdf = _make_pdf()
    via_filter, _ = extract_pdf(pdf, sections=["abstract"])
    doc = extract_sections(pdf)
    assert via_filter == (doc.abstract.text if doc.abstract else "")


def test_extract_docx_filter():
    pytest.importorskip("mammoth")
    pytest.importorskip("docx")
    from docx import Document
    d = Document()
    d.add_heading("Abstract", level=2); d.add_paragraph("Abstract body.")
    d.add_heading("Methods", level=2); d.add_paragraph("Methods body.")
    buf = io.BytesIO(); d.save(buf)
    from docpluck import extract_docx
    text, _ = extract_docx(buf.getvalue(), sections=["abstract"])
    assert "Abstract body" in text
    assert "Methods body" not in text


def test_extract_html_filter():
    from docpluck import extract_html
    html = b"<html><body><h2>Abstract</h2><p>Abstract body.</p>" \
           b"<h2>Methods</h2><p>Methods body.</p></body></html>"
    text, _ = extract_html(html, sections=["abstract"])
    assert "Abstract body" in text
    assert "Methods body" not in text
