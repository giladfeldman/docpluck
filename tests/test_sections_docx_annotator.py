"""DOCX annotator: heading detection via mammoth-emitted HTML."""

import io

import pytest

mammoth = pytest.importorskip("mammoth")
docx_pkg = pytest.importorskip("docx")  # python-docx for fixture creation

from docpluck.sections.annotators.docx import annotate_docx


def _build_docx_with_real_headings() -> bytes:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE  # noqa: F401
    d = Document()
    d.add_heading("Some Title", level=1)
    d.add_paragraph("Author, J.")
    d.add_heading("Abstract", level=2)
    d.add_paragraph("This paper investigates X.")
    d.add_heading("Methods", level=2)
    d.add_paragraph("We did things.")
    d.add_heading("References", level=2)
    d.add_paragraph("[1] Doe, J. (2020).")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_docx_with_real_headings():
    text, hints = annotate_docx(_build_docx_with_real_headings())
    assert isinstance(text, str)
    heading_texts = [h.text for h in hints if h.is_heading_candidate]
    assert "Abstract" in heading_texts
    assert "Methods" in heading_texts
    assert "References" in heading_texts
    for h in hints:
        if h.is_heading_candidate:
            assert h.heading_source == "markup"


def test_extract_sections_from_docx_bytes():
    from docpluck import extract_sections
    doc = extract_sections(_build_docx_with_real_headings())
    assert doc.source_format == "docx"
    assert doc.abstract is not None
    assert "investigates X" in doc.abstract.text
    assert doc.references is not None
