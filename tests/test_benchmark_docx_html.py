"""Benchmark tests for DOCX and HTML extraction.

These tests parallel the PDF benchmark methodology in docs/BENCHMARKS.md:
- Ground-truth passages must survive extraction (fuzzy match >= 90%)
- Extraction + normalization must be idempotent
- Quality scores on DOCX/HTML should match or exceed PDF scores
- Performance must stay within reasonable bounds

The ground-truth passages are synthesized into DOCX and HTML fixtures
programmatically so the suite runs anywhere without external test files.
The passages match real patterns from the PDF ground-truth set.
"""
import io
import time

import pytest

pytest.importorskip("mammoth", reason="mammoth not installed")
pytest.importorskip("bs4", reason="beautifulsoup4 not installed")
pytest.importorskip("lxml", reason="lxml not installed")
pytest.importorskip("docx", reason="python-docx not installed (dev dep)")

from docx import Document
from docpluck import (
    extract_docx,
    extract_html,
    normalize_text,
    NormalizationLevel,
    compute_quality_score,
)


# ---------------------------------------------------------------------------
# Ground-truth passages (mirror the PDF benchmark corpus)
# ---------------------------------------------------------------------------

# Each passage is something a downstream consumer (ESCIcheck, MetaESCI) needs
# to match with regex after extraction. These are the highest-stakes values.
GROUND_TRUTH_PASSAGES = [
    # Correlations with CIs
    "r(261) = -0.73, 95% CI [-0.78, -0.67], p < .001",
    "r(148) = 0.42, 95% CI [0.27, 0.55], p < .001",
    # F-tests
    "F(1, 30) = 4.42, p = .043",
    "F(2, 58) = 12.15, p < .001, eta2 = .30",
    # t-tests
    "t(98) = 2.15, p = .034, d = 0.43",
    "t(149) = -3.21, p = .002",
    # Chi-square
    "chi2(2, N = 250) = 8.42, p = .015",
    # Means and SDs
    "M = 3.42, SD = 0.89",
    "M = 4.21 (SD = 1.12)",
    # Regression
    "b = 0.34, SE = 0.08, t(197) = 4.25, p < .001",
    # Cohen's d
    "Cohen's d = 0.42, 95% CI [0.18, 0.66]",
    # Odds ratio
    "OR = 2.34, 95% CI [1.45, 3.78], p = .001",
    # Smart-quote variations
    '"significant main effect"',
    # Greek letters (via HTML entities in HTML, plain text in DOCX)
    "alpha = .05",
    # Negative value with Unicode minus
    "r = -0.42",
]


def _fuzzy_match(needle: str, haystack: str, threshold: int = 90) -> bool:
    """Use rapidfuzz partial_ratio matching, same threshold as PDF benchmark."""
    try:
        from rapidfuzz import fuzz
    except ImportError:
        # Fallback: substring match (stricter than fuzzy)
        return needle in haystack
    return fuzz.partial_ratio(needle, haystack) >= threshold


def _build_docx_fixture(passages: list[str]) -> bytes:
    """Create a realistic academic DOCX with the ground-truth passages embedded."""
    doc = Document()
    doc.add_heading("Study Results", level=1)
    doc.add_paragraph(
        "This section reports the statistical findings from the main analyses."
    )
    doc.add_heading("Primary Analyses", level=2)
    for passage in passages:
        doc.add_paragraph(f"We observed: {passage}.")
    doc.add_heading("Discussion", level=2)
    doc.add_paragraph("The results support our primary hypothesis.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_html_fixture(passages: list[str]) -> bytes:
    """Create a realistic publisher-style HTML with ground-truth passages."""
    body_items = "\n".join(
        f"<p>We observed: <em>{p}</em>.</p>" for p in passages
    )
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="utf-8">
    <title>Test Article</title>
    <style>body {{ font-family: sans-serif; }}</style>
    <script>var x = 1;</script>
</head>
<body>
    <article>
        <header>
            <h1>A Study of Things</h1>
            <p class="authors"><span>Chan</span><span>ORCID</span>, <span>Feldman</span></p>
        </header>
        <section>
            <h2>Study Results</h2>
            <p>This section reports the statistical findings from the main analyses.</p>
            <h3>Primary Analyses</h3>
            {body_items}
            <h3>Discussion</h3>
            <p>The results support our primary hypothesis.</p>
        </section>
    </article>
</body>
</html>"""
    return html.encode("utf-8")


# ---------------------------------------------------------------------------
# Ground-truth accuracy
# ---------------------------------------------------------------------------

class TestDocxGroundTruth:
    def test_all_passages_survive_extraction(self):
        docx = _build_docx_fixture(GROUND_TRUTH_PASSAGES)
        text, method = extract_docx(docx)
        assert method == "mammoth"

        missing = []
        for passage in GROUND_TRUTH_PASSAGES:
            if not _fuzzy_match(passage, text):
                missing.append(passage)

        assert not missing, f"Missing passages: {missing}"

    def test_academic_normalization_preserves_passages(self):
        docx = _build_docx_fixture(GROUND_TRUTH_PASSAGES)
        text, _ = extract_docx(docx)
        normalized, _ = normalize_text(text, NormalizationLevel.academic)

        missing = []
        for passage in GROUND_TRUTH_PASSAGES:
            if not _fuzzy_match(passage, normalized):
                missing.append(passage)

        assert not missing, f"Missing after normalization: {missing}"


class TestHtmlGroundTruth:
    def test_all_passages_survive_extraction(self):
        html = _build_html_fixture(GROUND_TRUTH_PASSAGES)
        text, method = extract_html(html)
        assert method == "beautifulsoup"

        missing = []
        for passage in GROUND_TRUTH_PASSAGES:
            if not _fuzzy_match(passage, text):
                missing.append(passage)

        assert not missing, f"Missing passages: {missing}"

    def test_chan_orcid_regression_in_realistic_fixture(self):
        """The real-world bug that went undetected for weeks in Scimeto."""
        html = _build_html_fixture(GROUND_TRUTH_PASSAGES)
        text, _ = extract_html(html)
        assert "ChanORCID" not in text
        assert "Chan" in text
        assert "ORCID" in text

    def test_script_and_style_stripped(self):
        html = _build_html_fixture(GROUND_TRUTH_PASSAGES)
        text, _ = extract_html(html)
        assert "var x = 1" not in text
        assert "font-family" not in text

    def test_academic_normalization_preserves_passages(self):
        html = _build_html_fixture(GROUND_TRUTH_PASSAGES)
        text, _ = extract_html(html)
        normalized, _ = normalize_text(text, NormalizationLevel.academic)

        missing = []
        for passage in GROUND_TRUTH_PASSAGES:
            if not _fuzzy_match(passage, normalized):
                missing.append(passage)

        assert not missing, f"Missing after normalization: {missing}"


# ---------------------------------------------------------------------------
# Idempotency
# ---------------------------------------------------------------------------

class TestIdempotency:
    def test_docx_normalization_idempotent(self):
        docx = _build_docx_fixture(GROUND_TRUTH_PASSAGES)
        text, _ = extract_docx(docx)
        text1, _ = normalize_text(text, NormalizationLevel.academic)
        text2, _ = normalize_text(text1, NormalizationLevel.academic)
        assert text1 == text2

    def test_html_normalization_idempotent(self):
        html = _build_html_fixture(GROUND_TRUTH_PASSAGES)
        text, _ = extract_html(html)
        text1, _ = normalize_text(text, NormalizationLevel.academic)
        text2, _ = normalize_text(text1, NormalizationLevel.academic)
        assert text1 == text2


# ---------------------------------------------------------------------------
# Quality scoring
# ---------------------------------------------------------------------------

class TestQualityScores:
    def test_docx_quality_is_high(self):
        docx = _build_docx_fixture(GROUND_TRUTH_PASSAGES)
        text, _ = extract_docx(docx)
        quality = compute_quality_score(text)
        # DOCX should score at least as well as PDF (fewer extraction artifacts)
        assert quality["score"] >= 80
        assert quality["confidence"] == "high"
        assert quality["details"]["garbled_chars"] == 0

    def test_html_quality_is_high(self):
        html = _build_html_fixture(GROUND_TRUTH_PASSAGES)
        text, _ = extract_html(html)
        quality = compute_quality_score(text)
        assert quality["score"] >= 80
        assert quality["confidence"] == "high"
        assert quality["details"]["garbled_chars"] == 0


# ---------------------------------------------------------------------------
# Performance
# ---------------------------------------------------------------------------

class TestPerformance:
    def test_docx_extraction_under_1s(self):
        docx = _build_docx_fixture(GROUND_TRUTH_PASSAGES)
        start = time.perf_counter()
        extract_docx(docx)
        elapsed = time.perf_counter() - start
        # Should be well under 1 second for a small fixture
        assert elapsed < 1.0, f"DOCX extraction took {elapsed:.2f}s"

    def test_html_extraction_under_1s(self):
        html = _build_html_fixture(GROUND_TRUTH_PASSAGES)
        start = time.perf_counter()
        extract_html(html)
        elapsed = time.perf_counter() - start
        assert elapsed < 1.0, f"HTML extraction took {elapsed:.2f}s"
